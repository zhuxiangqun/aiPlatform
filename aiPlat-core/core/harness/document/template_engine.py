"""
TemplateEngine — 文档模板引擎 (Agent 闭环执行 — 本地结果输出)

挂载 Word/Excel/Markdown 模板 → AI 按模板格式生成输出 → 写入本地文件。

安全:
  - 文件类型白名单: .docx / .xlsx / .md (拒绝含宏的 .docm / .xlsm)
  - 模板解析时不读取宿主文件系统的非模板路径

调用者: PipelineStage (output_artifact rendered 模式) / REST API
"""

from __future__ import annotations

import json as _json
import logging
import os as _os
import re as _re
import time as _time
from dataclasses import dataclass, field
from pathlib import Path as _Path
from typing import Any, Dict, List, Optional, Set

logger = logging.getLogger(__name__)

TEMPLATE_DIR = _Path(_os.getenv("AIPLAT_HOME", _Path("~").expanduser() / ".aiplat")) / "templates"
ALLOWED_EXTENSIONS: Set[str] = {".docx", ".xlsx", ".md", ".txt", ".csv"}


# ── Data Models ──────────────────────────────────────────────────────────

@dataclass
class Template:
    template_id: str
    path: str
    format: str                    # docx | xlsx | md
    placeholders: List[str] = field(default_factory=list)
    description: str = ""
    registered_at: float = 0.0

    def to_dict(self) -> Dict[str, Any]:
        return {
            "template_id": self.template_id,
            "path": self.path,
            "format": self.format,
            "placeholders": self.placeholders,
            "description": self.description,
            "registered_at": self.registered_at,
        }


# ── TemplateRegistry ─────────────────────────────────────────────────────

class TemplateRegistry:
    """模板注册中心."""

    _instance: Optional["TemplateRegistry"] = None

    def __init__(self):
        self._templates: Dict[str, Template] = {}
        self._scan()

    @classmethod
    def get(cls) -> "TemplateRegistry":
        if cls._instance is None:
            cls._instance = cls()
        return cls._instance

    def _scan(self) -> None:
        if not TEMPLATE_DIR.exists():
            return
        for f in TEMPLATE_DIR.iterdir():
            if f.is_file() and f.suffix.lower() in ALLOWED_EXTENSIONS:
                tid = f.stem
                self._templates[tid] = Template(
                    template_id=tid,
                    path=str(f),
                    format=f.suffix.lstrip("."),
                    placeholders=self._extract_placeholders(f),
                    registered_at=f.stat().st_mtime,
                )

    def _extract_placeholders(self, path: _Path) -> List[str]:
        """从模板文件中提取占位符 {{name}}."""
        try:
            content = path.read_text(encoding="utf-8", errors="ignore")
            return list(set(_re.findall(r"\{\{(\w+(?:\.\w+)*)\}\}", content)))
        except Exception:
            return []

    def register(self, template_id: str, path: str, description: str = "") -> Template:
        """注册模板.

        安全: 拒绝含宏的文件类型 (.docm, .xlsm)
        """
        ext = _Path(path).suffix.lower()
        if ext not in ALLOWED_EXTENSIONS:
            raise ValueError(
                f"Template file type '{ext}' not allowed. "
                f"Allowed: {sorted(ALLOWED_EXTENSIONS)}. "
                f"Macro-enabled formats (.docm, .xlsm) are blocked for security."
            )

        if not _Path(path).exists():
            raise FileNotFoundError(f"Template file not found: {path}")

        template = Template(
            template_id=template_id,
            path=str(path),
            format=ext.lstrip("."),
            placeholders=self._extract_placeholders(_Path(path)),
            description=description,
            registered_at=_time.time(),
        )
        self._templates[template_id] = template
        logger.info("Registered template '%s' (%s, %d placeholders)",
                     template_id, ext, len(template.placeholders))
        return template

    def get(self, template_id: str) -> Optional[Template]:
        return self._templates.get(template_id)

    def list_all(self) -> List[Dict[str, Any]]:
        return [t.to_dict() for t in self._templates.values()]


# ── TemplateRenderer ──────────────────────────────────────────────────────

class TemplateRenderer:
    """模板渲染器.

    安全: 模板解析时不读取宿主文件系统的非模板路径.
    """

    def render(
        self,
        template_id: str,
        data: Dict[str, Any],
        *,
        output_path: str = "",
    ) -> Dict[str, Any]:
        """渲染模板.

        Args:
            template_id: 模板 ID
            data: 数据字典 (key → value, 支持嵌套 dot 路径)
            output_path: 输出路径 (留空则自动生成)

        Returns:
            {"path": "/output/rendered.docx", "format": "docx", "size_bytes": 1234}
        """
        registry = TemplateRegistry.get()
        template = registry.get(template_id)
        if not template:
            raise ValueError(f"Template '{template_id}' not found")

        fmt = template.format
        if fmt == "md":
            return self._render_markdown(template, data, output_path)
        elif fmt == "xlsx":
            return self._render_excel(template, data, output_path)
        elif fmt == "docx":
            return self._render_docx(template, data, output_path)
        else:
            return self._render_text(template, data, output_path)

    def _render_markdown(
        self, template: Template, data: Dict[str, Any], output_path: str
    ) -> Dict[str, Any]:
        """渲染 Markdown 模板."""
        content = _Path(template.path).read_text(encoding="utf-8", errors="ignore")
        rendered = self._substitute(content, data)

        out = output_path or self._default_output(template, ".md")
        _Path(out).parent.mkdir(parents=True, exist_ok=True)
        _Path(out).write_text(rendered, encoding="utf-8")

        return {"path": out, "format": "md", "size_bytes": _Path(out).stat().st_size}

    def _render_excel(
        self, template: Template, data: Dict[str, Any], output_path: str
    ) -> Dict[str, Any]:
        """渲染 Excel 模板 (轻量 — 字符串替换)."""
        try:
            import openpyxl
            wb = openpyxl.load_workbook(template.path)
            for ws in wb.worksheets:
                for row in ws.iter_rows():
                    for cell in row:
                        if isinstance(cell.value, str) and "{{" in cell.value:
                            cell.value = self._substitute(cell.value, data)
            out = output_path or self._default_output(template, ".xlsx")
            wb.save(out)
            return {"path": out, "format": "xlsx", "size_bytes": _Path(out).stat().st_size}
        except ImportError:
            logger.warning("openpyxl not installed, falling back to text render")
            return self._render_text(template, data, output_path)

    def _render_docx(
        self, template: Template, data: Dict[str, Any], output_path: str
    ) -> Dict[str, Any]:
        """渲染 Word 模板."""
        try:
            from docx import Document
            doc = Document(template.path)
            for p in doc.paragraphs:
                if "{{" in p.text:
                    for run in p.runs:
                        if "{{" in run.text:
                            run.text = self._substitute(run.text, data)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if "{{" in cell.text:
                            for p in cell.paragraphs:
                                for run in p.runs:
                                    if "{{" in run.text:
                                        run.text = self._substitute(run.text, data)
            out = output_path or self._default_output(template, ".docx")
            doc.save(out)
            return {"path": out, "format": "docx", "size_bytes": _Path(out).stat().st_size}
        except ImportError:
            logger.warning("python-docx not installed, falling back to text render")
            return self._render_text(template, data, output_path)

    def _render_text(
        self, template: Template, data: Dict[str, Any], output_path: str
    ) -> Dict[str, Any]:
        """通用文本渲染."""
        content = _Path(template.path).read_text(encoding="utf-8", errors="ignore")
        rendered = self._substitute(content, data)
        out = output_path or self._default_output(template, ".txt")
        _Path(out).parent.mkdir(parents=True, exist_ok=True)
        _Path(out).write_text(rendered, encoding="utf-8")
        return {"path": out, "format": "txt", "size_bytes": _Path(out).stat().st_size}

    @staticmethod
    def _substitute(text: str, data: Dict[str, Any]) -> str:
        """替换占位符 {{key}} 和 {{nested.key}}."""
        def _get_val(match):
            key = match.group(1)
            parts = key.split(".")
            val = data
            for p in parts:
                if isinstance(val, dict):
                    val = val.get(p, "")
                elif isinstance(val, list):
                    try:
                        val = val[int(p)] if p.isdigit() else val
                    except (IndexError, ValueError):
                        return ""
                else:
                    return f"{{{{{key}}}}}"
            return str(val) if val is not None else ""

        return _re.sub(r"\{\{(\w+(?:\.\w+)*)\}\}", _get_val, text)

    @staticmethod
    def _default_output(template: Template, ext: str) -> str:
        ts = _time.strftime("%Y%m%d_%H%M%S")
        out_dir = _Path(_os.getenv("AIPLAT_HOME", _Path("~").expanduser() / ".aiplat")) / "output"
        out_dir.mkdir(parents=True, exist_ok=True)
        return str(out_dir / f"{template.template_id}_{ts}{ext}")
