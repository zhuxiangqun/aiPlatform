"""DOCX converter — delegates to MarkItDown with python-docx fallback."""
import os
import re
from typing import Any, BinaryIO, List

from core.harness.document.protocol import (
    DocumentConverter, DocumentElement, StreamInfo, detect_structure_role,
)


class DocxConverter(DocumentConverter):
    """DOCX → Markdown via MarkItDown (or python-docx fallback)."""

    SOURCE_FORMAT = "docx"
    REQUIRED_PACKAGES = {}  # markitdown + python-docx are soft deps
    ACCEPTED_EXTENSIONS = (".docx", ".doc")
    ACCEPTED_MIME_PREFIXES = (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "application/msword",
    )

    def accepts(
        self,
        file_stream: BinaryIO,
        stream_info: StreamInfo,
        **kwargs: Any,
    ) -> bool:
        return self._accepts_by_format(stream_info)

    def convert(
        self,
        file_stream: BinaryIO,
        stream_info: StreamInfo,
        **kwargs: Any,
    ) -> List[DocumentElement]:
        file_path = stream_info.local_path
        if not file_path:
            return self._convert_stream_docx(file_stream)

        try:
            from markitdown import MarkItDown
            return self._convert_via_markitdown(file_path)
        except ImportError:
            return self._fallback_docx(file_path)

    def _fallback_docx(self, file_path: str) -> List[DocumentElement]:
        try:
            from docx import Document
        except ImportError:
            return self._fallback_text_from_file(file_path)

        doc = Document(file_path)
        elements: List[DocumentElement] = []
        for para in doc.paragraphs:
            text = (para.text or "").strip()
            if not text:
                continue
            elements.append(DocumentElement(
                type="text", text=text, page_idx=0,
                meta={"source": "docx"},
                source_format="docx",
            ))
        for ti, table in enumerate(doc.tables):
            rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
            if any(any(c for c in row) for row in rows):
                elements.append(DocumentElement(
                    type="table",
                    text="\n".join(" | ".join(r) for r in rows),
                    page_idx=0, cells=rows,
                    meta={"source": "docx", "table_index": ti},
                    source_format="docx",
                ))
        return elements or self._fallback_text_from_file(file_path)

    def _convert_stream_docx(self, file_stream: BinaryIO) -> List[DocumentElement]:
        try:
            from docx import Document
            doc = Document(file_stream)
        except ImportError:
            data = file_stream.read()
            text = data.decode("utf-8", errors="ignore")
            return [DocumentElement(
                type="text", text=text.strip(), page_idx=0,
                meta={"source": "docx", "fallback": True},
                source_format="docx",
            )] if text.strip() else []

        elements: List[DocumentElement] = []
        for para in doc.paragraphs:
            text = (para.text or "").strip()
            if not text:
                continue
            elements.append(DocumentElement(
                type="text", text=text, page_idx=0,
                meta={"source": "docx"},
                source_format="docx",
            ))
        for ti, table in enumerate(doc.tables):
            rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
            if any(any(c for c in row) for row in rows):
                elements.append(DocumentElement(
                    type="table",
                    text="\n".join(" | ".join(r) for r in rows),
                    page_idx=0, cells=rows,
                    meta={"source": "docx", "table_index": ti},
                    source_format="docx",
                ))
        return elements
