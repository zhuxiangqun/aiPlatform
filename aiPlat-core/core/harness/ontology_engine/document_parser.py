"""
Document Parser — 多格式文档解析器。

支持格式: Markdown (.md), HTML (.html), Plain Text (.txt)
软依赖: PDF (.pdf) 需 PyPDF2/pdfplumber, Word (.docx) 需 python-docx

输出: StructuredChunk[] — 含章节路径、文本、NER实体候选
"""

from __future__ import annotations

import re as _re
from dataclasses import dataclass, field
from pathlib import Path as _Path
from typing import Any, Dict, List, Optional


@dataclass
class StructuredChunk:
    """A parsed document chunk with structural metadata."""
    id: str
    text: str
    heading_path: List[str] = field(default_factory=list)
    page_num: int = 0
    chunk_index: int = 0
    entities: List[Dict[str, Any]] = field(default_factory=list)
    metadata: Dict[str, Any] = field(default_factory=dict)

    def to_dict(self) -> dict:
        return {
            "id": self.id,
            "text": self.text,
            "heading_path": self.heading_path,
            "page_num": self.page_num,
            "chunk_index": self.chunk_index,
            "entities": self.entities,
            "metadata": self.metadata,
        }


@dataclass
class StructuredTable:
    """A structured table extracted from a document, preserving row/column relationships.

    Unlike flattened text extraction, this retains the relational structure:
    each row is a record, each column is a field, and headers map to values.
    This enables the PropertyExtractor and GraphIndex to work with structured
    data instead of broken text fragments.
    """
    table_id: str                              # unique identifier (e.g. "chunk-0-table-0")
    caption: str = ""                          # table title/caption
    headers: List[str] = field(default_factory=list)  # column names
    rows: List[List[str]] = field(default_factory=list)  # data rows
    page_num: int = 0
    chunk_id: str = ""                         # parent chunk

    def to_dict(self) -> dict:
        return {
            "table_id": self.table_id,
            "caption": self.caption,
            "headers": self.headers,
            "rows": self.rows[:20],  # Truncate for display
            "row_count": len(self.rows),
            "page_num": self.page_num,
            "chunk_id": self.chunk_id,
        }

    def to_text(self) -> str:
        """Render table as readable text for LLM consumption."""
        lines = []
        if self.caption:
            lines.append(f"表格: {self.caption}")
        if self.headers:
            lines.append(" | ".join(self.headers))
            lines.append(" | ".join("---" for _ in self.headers))
        for row in self.rows:
            lines.append(" | ".join(str(c) for c in row))
        return "\n".join(lines)

    def to_entity_rows(self) -> List[Dict[str, str]]:
        """Convert each row to a dict of {header → value}."""
        if not self.headers:
            return []
        result = []
        for row in self.rows:
            record = {}
            for i, h in enumerate(self.headers):
                record[h] = row[i] if i < len(row) else ""
            result.append(record)
        return result


@dataclass
class ParsedDocument:
    """Complete parsed document with metadata and chunks."""
    title: str = ""
    format: str = ""
    raw_text: str = ""
    chunks: List[StructuredChunk] = field(default_factory=list)
    tables: List[StructuredTable] = field(default_factory=list)
    metadata: Dict[str, Any] = field(default_factory=dict)
    parse_warnings: List[str] = field(default_factory=list)

    def to_dict(self) -> dict:
        return {
            "title": self.title,
            "format": self.format,
            "raw_text": self.raw_text[:2000],
            "chunks": [c.to_dict() for c in self.chunks],
            "chunk_count": len(self.chunks),
            "tables": [t.to_dict() for t in self.tables],
            "table_count": len(self.tables),
            "metadata": self.metadata,
            "warnings": self.parse_warnings,
        }


@dataclass
class QAPair:
    """A question-answer pair extracted from FAQ-style documents.

    K2 Structure Planning: FAQ, expert experience, and customer service scripts
    should be organized as QA pairs, not just flat text chunks. This preserves
    the question→answer relationship for precise retrieval.
    """
    id: str
    question: str
    answer: str
    tags: List[str] = field(default_factory=list)
    source_chunk_id: str = ""

    def to_dict(self) -> dict:
        return {
            "id": self.id, "question": self.question, "answer": self.answer[:500],
            "tags": self.tags, "source_chunk_id": self.source_chunk_id,
        }


@dataclass
class ImageSegment:
    """Multi-modal storage: image region with context for retrieval."""
    id: str
    page_num: int = 0
    description: str = ""            # VLM-generated natural language description
    image_base64: str = ""           # cropped image for display
    bbox: Dict[str, int] = field(default_factory=dict)  # {x, y, w, h}
    confidence: float = 0.0
    figure_type: str = ""            # chart, diagram, photo, table_image
    transcription: str = ""          # axis labels, legend text


class AdaptiveRouter:
    """Route pages to optimal parser: OCR for text-dense, VLM for image-dense."""

    @staticmethod
    def route(page_text: str, image_regions: int = 0) -> str:
        """Returns 'ocr' | 'vlm' based on text density heuristic."""
        text_chars = len(page_text.strip()) if page_text else 0
        # If very little text but has images → VLM
        if text_chars < 100 and image_regions > 0:
            return "vlm"
        # If text density < 30% of expected page → VLM
        if text_chars < 500 and image_regions >= 3:
            return "vlm"
        return "ocr"


class VisualPageParser:
    """VLM-based parser for chart/image-heavy pages. Soft dependency on LLM adapter."""

    @staticmethod
    async def parse_page(
        image_base64: str,
        *,
        model_name: str = "",
        page_num: int = 0,
    ) -> Dict[str, Any]:
        """Send page image to VLM, return structured output with figure descriptions.

        Returns: {markdown, figures: [{kind, description, transcription, confidence}]}
        """
        prompt = (
            "Extract all content from this document page image. Return JSON:\n"
            '{\n  "markdown": "full page text in markdown",\n'
            '  "figures": [\n'
            '    {\n'
            '      "kind": "line_chart|bar_chart|flowchart|diagram|photo|table_image",\n'
            '      "description": "natural language description (50-150 chars)",\n'
            '      "transcription": "axis labels, legend text, readable values",\n'
            '      "confidence": 0.8\n'
            '    }\n'
            '  ]\n'
            '}\n\n'
            'If no figures found, return empty figures list. Only return JSON.'
        )

        try:
            from core.harness.syscalls.llm import sys_llm_generate
            from core.harness.utils.model_injection import best_model_for_purpose

            resp = await sys_llm_generate(
                None,
                [
                    {"role": "system", "content": "You are a document parsing expert."},
                    {"role": "user", "content": [
                        {"type": "text", "text": prompt},
                        {"type": "image_url", "image_url": {"url": f"data:image/png;base64,{image_base64}"}},
                    ]},
                ],
                model_name=model_name or best_model_for_purpose("chat"),
                temperature=0.1,
                max_tokens=2000,
            )
            import json as _json
            import re as _re
            text = getattr(resp, 'content', '') or str(resp)
            # Extract JSON from response
            m = _re.search(r'\{[\s\S]*\}', text)
            if m:
                return _json.loads(m.group(0))
            return {"markdown": "", "figures": [], "error": "no_json"}
        except Exception as e:
            return {"markdown": "", "figures": [], "error": str(e)}

    @staticmethod
    def check_completeness(
        expected_figures: int,
        parsed_figures: int,
    ) -> Dict[str, Any]:
        """Chart completeness check. Returns {complete, missing, action}."""
        missing = expected_figures - parsed_figures
        return {
            "complete": missing <= 0,
            "expected": expected_figures,
            "parsed": parsed_figures,
            "missing": max(0, missing),
            "action": "retry_vlm" if missing > 2 else "ok" if missing <= 0 else "review",
        }


class DocumentParser:
    """Multi-format document parser.
    
    Delegates raw text extraction to ConverterRegistry (System A),
    and adds structural enrichment: heading paths, structured tables,
    QA pairs, entity candidates, VLM page parsing.
    
    This implements the "API entry uniqueness" principle (§5.7)
    — all raw parsing converges through the ConverterRegistry.
    """

    # ── Public API ──

    def parse_file(self, file_path: str) -> ParsedDocument:
        """Parse a document file by path. Auto-detects format from extension.
        
        Supports: pdf, docx, html, md, txt, mp4, avi, mov, wav, mp3, png, jpg
        
        Now delegates dispatch to ConverterRegistry (System A) for raw text extraction,
        adding System B's structural enrichment on top.
        """
        from core.api.facades.kb_facade import _KIND_TO_EXT, normalize_kind
        path = _Path(file_path)
        ext = path.suffix.lower()
        kind = normalize_kind(ext.lstrip("."))
        
        # Delegate to registry for formats that have converters
        if kind in ("pdf", "docx", "pptx", "xlsx", "video", "audio"):
            if kind == "pdf":
                return self._parse_pdf(path.read_bytes(), file_path)
            elif kind == "docx":
                return self._parse_docx(path.read_bytes(), file_path)
            elif kind == "video":
                return self._parse_video(file_path)
            elif kind == "audio":
                return self._parse_audio_file(file_path)
            elif kind == "image":
                return self._parse_image_file(file_path)
            else:
                # Other registry formats: pass through raw text
                elements = self._parse_via_registry(file_path, ext)
                if elements:
                    return self._parse_text("\n\n".join(
                        el.text for el in elements if el.text.strip()
                    ), file_path)
                return self._parse_text("", file_path)
        elif kind == "html" or ext in (".html", ".htm"):
            return self._parse_html(path.read_bytes().decode("utf-8", errors="ignore"), file_path)
        elif kind in ("markdown", "md") or ext == ".md":
            return self._parse_markdown(path.read_bytes().decode("utf-8", errors="ignore"), file_path)
        elif kind == "image":
            return self._parse_image_file(file_path)
        else:
            return self._parse_text(path.read_bytes().decode("utf-8", errors="ignore"), file_path)

    def parse_text(self, text: str, *, title: str = "", format: str = "txt") -> ParsedDocument:
        """Parse raw text string. Uses canonical format normalization."""
        from core.api.facades.kb_facade import normalize_kind
        fmt = normalize_kind(format)
        if fmt == "markdown":
            return self._parse_markdown(text, title)
        elif fmt == "html":
            return self._parse_html(text, title)
        else:
            return self._parse_text(text, title)

    # ── Format Parsers ──

    def _parse_markdown(self, text: str, source: str = "") -> ParsedDocument:
        """Parse Markdown into heading-structured chunks."""
        doc = ParsedDocument(format="md", raw_text=text)
        doc.title = self._extract_title_md(text) or source
        chunks = self._split_by_headings(text)
        doc.chunks = self._build_chunks(chunks)
        return doc

    def _parse_html(self, text: str, source: str = "") -> ParsedDocument:
        """Parse HTML into structured chunks."""
        doc = ParsedDocument(format="html", raw_text=text)
        try:
            from bs4 import BeautifulSoup
            soup = BeautifulSoup(text, "html.parser")
            # Extract title
            title_tag = soup.find("title") or soup.find("h1")
            doc.title = title_tag.get_text().strip() if title_tag else source
            # Remove script/style
            for tag in soup(["script", "style", "nav", "footer", "header"]):
                tag.decompose()
            # Extract heading-structured content
            body = soup.find("body") or soup
            chunks = self._split_html_by_headings(body)
            doc.chunks = self._build_chunks(chunks)
        except Exception as e:
            doc.parse_warnings.append(f"HTML parse warning: {e}")
            # Fallback to plain text
            clean = _re.sub(r"<[^>]+>", " ", text)
            clean = _re.sub(r"\s+", " ", clean).strip()
            doc.title = source
            doc.chunks = self._build_chunks([(clean, ["正文"])])
        return doc

    def _clean_text(self, text: str) -> tuple:
        """Apply configurable cleanup rules to raw extracted text."""
        try:
            from core.harness.knowledge.text_cleaner import get_text_cleaner
            cleaner = get_text_cleaner()
            return cleaner.clean(text)
        except Exception:
            return text, 0

    def _parse_text(self, text: str, source: str = "") -> ParsedDocument:
        """Parse plain text into paragraph chunks."""
        # ── Apply configurable text cleanup ──
        text, removed = self._clean_text(text)
        if removed > 0:
            import logging
            logging.getLogger("document_parser").debug("Text cleanup: %d matches removed", removed)

        doc = ParsedDocument(format="txt", raw_text=text)
        # Extract first line as title
        lines = text.strip().split("\n")
        doc.title = lines[0].strip()[:120] if lines else source
        # Chunk by paragraph (double newline)
        paragraphs = _re.split(r"\n\s*\n", text)
        chunks = []
        for i, para in enumerate(paragraphs):
            para = para.strip()
            if len(para) > 15:
                chunks.append((para, []))
        doc.chunks = self._build_chunks(chunks)
        return doc

    def _parse_pdf(self, content: bytes, source: str = "") -> ParsedDocument:
        """Parse PDF via ConverterRegistry, with table extraction from pdfplumber.
        
        Text extraction delegates to the ConverterRegistry (System A).
        Table extraction is unique to System B — pdfplumber preserves
        row/column relationships that the registry converters flatten.
        """
        doc = ParsedDocument(format="pdf", raw_text="")
        doc.title = source

        # Step 1: Try registry-based parsing (markitdown or pdfplumber)
        elements = self._parse_via_registry(source, ".pdf")
        if elements:
            text = "\n\n".join(el.text for el in elements if el.text.strip())
            doc.raw_text = text
            doc.chunks = self._build_chunks_from_elements(elements)
        else:
            text = ""

        # Step 2: Extract tables with pdfplumber (System B unique value)
        all_tables: List[StructuredTable] = self._extract_pdf_tables(content)

        # Step 3: Fallback to pdfplumber/PyPDF2 if registry produced no text
        if not text:
            text = self._fallback_pdf_text(content, doc)
            if text:
                doc.raw_text = text
                doc.chunks = self._build_chunks_from_raw(text)

        if text and not doc.chunks:
            result = self._parse_text(text, source)
            result.tables = all_tables
            result.metadata["table_count"] = len(all_tables)
            self._attach_tables_to_chunks(result)
            return result

        doc.tables = all_tables
        doc.metadata["table_count"] = len(all_tables)
        if doc.chunks:
            self._attach_tables_to_chunks(doc)
        if not text and not doc.chunks:
            doc.parse_warnings.append("No text extracted from PDF")
        return doc

    def _extract_pdf_tables(self, content: bytes) -> List[StructuredTable]:
        """Extract tables from PDF content using pdfplumber."""
        tables: List[StructuredTable] = []
        try:
            import pdfplumber
            import io
            with pdfplumber.open(io.BytesIO(content)) as pdf:
                for i, page in enumerate(pdf.pages):
                    raw_tables = page.extract_tables()
                    if raw_tables:
                        for ti, raw in enumerate(raw_tables):
                            if raw and len(raw) > 1:
                                headers = [str(c or "") for c in raw[0]]
                                rows = [[str(c or "") for c in row] for row in raw[1:]]
                                tables.append(StructuredTable(
                                    table_id=f"pdf-page{i}-table{ti}",
                                    headers=headers,
                                    rows=rows,
                                    page_num=i,
                                ))
        except ImportError:
            pass
        except Exception as e:
            pass  # Table extraction is best-effort
        return tables

    def _fallback_pdf_text(self, content: bytes, doc: ParsedDocument) -> str:
        """Fallback PDF text extraction using pdfplumber or PyPDF2."""
        text = ""
        try:
            import pdfplumber
            import io
            with pdfplumber.open(io.BytesIO(content)) as pdf:
                pages = []
                for page in pdf.pages:
                    t = page.extract_text() or ""
                    if t.strip():
                        pages.append(t)
                text = "\n\n".join(pages)
                doc.metadata["page_count"] = len(pdf.pages)
        except ImportError:
            pass
        except Exception as e:
            doc.parse_warnings.append(f"pdfplumber error: {e}")

        if not text:
            try:
                import PyPDF2
                import io
                reader = PyPDF2.PdfReader(io.BytesIO(content))
                pages = []
                for page in reader.pages:
                    t = page.extract_text() or ""
                    if t.strip():
                        pages.append(t)
                text = "\n\n".join(pages)
                doc.metadata["page_count"] = len(reader.pages)
            except ImportError:
                doc.parse_warnings.append(
                    "PDF parsing requires pdfplumber or PyPDF2. Install: pip install pdfplumber"
                )
            except Exception as e:
                doc.parse_warnings.append(f"PyPDF2 error: {e}")
        return text

    def _parse_docx(self, content: bytes, source: str = "") -> ParsedDocument:
        """Parse Word document via ConverterRegistry + table extraction."""
        doc = ParsedDocument(format="docx", raw_text="")

        # Step 1: Try registry-based parsing (markitdown)
        elements = self._parse_via_registry(source, ".docx")
        if elements:
            doc.raw_text = "\n\n".join(el.text for el in elements if el.text.strip())
            doc.chunks = self._build_chunks_from_elements(elements)
            doc.title = doc.chunks[0].text[:120] if doc.chunks else source
        else:
            doc.title = source

        # Step 2: Extract tables via python-docx (System B unique value)
        all_tables: List[StructuredTable] = []
        try:
            import docx as _docx
            import io
            d = _docx.Document(io.BytesIO(content))
            for ti, table in enumerate(d.tables):
                rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
                if rows and len(rows) > 1:
                    headers = rows[0]
                    data_rows = rows[1:]
                    all_tables.append(StructuredTable(
                        table_id=f"docx-table{ti}",
                        headers=headers,
                        rows=data_rows,
                    ))
        except ImportError:
            pass
        except Exception as e:
            doc.parse_warnings.append(f"DOCX table extraction error: {e}")

        doc.tables = all_tables
        doc.metadata["table_count"] = len(all_tables)

        # Step 3: Fallback to python-docx for text if registry failed
        if not doc.chunks:
            try:
                import docx as _docx
                import io
                d = _docx.Document(io.BytesIO(content))
                paragraphs = []
                for p in d.paragraphs:
                    style = p.style.name if p.style else ""
                    text = p.text.strip()
                    if text:
                        level = 0
                        if "Heading" in style or "heading" in style:
                            try:
                                level = int(_re.search(r"\d+", style).group())
                            except Exception:
                                level = 1
                        paragraphs.append((text, level))
                doc.title = paragraphs[0][0][:120] if paragraphs else source
                chunks = self._build_heading_chunks(paragraphs)
                doc.chunks = self._build_chunks(chunks)
            except ImportError:
                doc.parse_warnings.append(
                    "Word parsing requires python-docx. Install: pip install python-docx"
                )
            except Exception as e:
                doc.parse_warnings.append(f"DOCX parse error: {e}")

        if doc.chunks:
            self._attach_tables_to_chunks(doc)
        elif not doc.chunks and not doc.parse_warnings:
            text = content.decode("utf-8", errors="ignore")
            return self._parse_text(text, source)
        return doc

    # ── Chunking ──

    def _split_by_headings(self, text: str) -> List[tuple]:
        """Split Markdown by heading levels.

        Uses (heading_text, level) tuples to correctly handle skipped
        heading levels (e.g. ### without preceding # or ##).
        """
        lines = text.split("\n")
        chunks = []
        current_stack: list = []  # [(heading_text, level), ...]
        current_lines: list = []

        for line in lines:
            h_match = _re.match(r"^(#{1,6})\s+(.+)", line)
            if h_match:
                if current_lines:
                    heading_path = [h for h, _ in current_stack]
                    chunks.append(("\n".join(current_lines), heading_path))
                    current_lines = []
                level = len(h_match.group(1))
                heading = h_match.group(2).strip()
                # Remove all headings at this level or deeper
                current_stack = [(h, l) for h, l in current_stack if l < level]
                current_stack.append((heading, level))
            else:
                current_lines.append(line)

        if current_lines:
            heading_path = [h for h, _ in current_stack]
            chunks.append(("\n".join(current_lines), heading_path))

        return chunks

    def _split_html_by_headings(self, body) -> List[tuple]:
        """Split HTML by heading tags. Uses (heading_text, level) tuples."""
        try:
            from bs4 import Tag, NavigableString
        except ImportError:
            return [(" ".join(body.stripped_strings)[:5000], [])]

        chunks = []
        current_stack: list = []  # [(heading_text, level), ...]
        current_text: list = []

        for elem in body.descendants:
            if isinstance(elem, Tag) and elem.name in ("h1", "h2", "h3", "h4", "h5", "h6"):
                if current_text:
                    heading_path = [h for h, _ in current_stack]
                    chunks.append((" ".join(current_text), heading_path))
                    current_text = []
                level = int(elem.name[1])
                h_text = elem.get_text().strip()
                current_stack = [(h, l) for h, l in current_stack if l < level]
                current_stack.append((h_text, level))
            elif isinstance(elem, NavigableString):
                t = str(elem).strip()
                if t and len(t) > 5:
                    current_text.append(t)

        if current_text:
            heading_path = [h for h, _ in current_stack]
            chunks.append((" ".join(current_text), heading_path))

        return chunks

    def _build_heading_chunks(self, paragraphs: List[tuple]) -> List[tuple]:
        """Build chunks from (text, level) tuples."""
        chunks = []
        current_heading = []
        current_lines = []
        for text, level in paragraphs:
            if level > 0:
                if current_lines:
                    chunks.append(("\n".join(current_lines), list(current_heading)))
                    current_lines = []
                current_heading = current_heading[: level - 1]
                current_heading.append(text)
            else:
                current_lines.append(text)
        if current_lines:
            chunks.append(("\n".join(current_lines), list(current_heading)))
        return chunks

    def _build_chunks(self, chunks: List[tuple]) -> List[StructuredChunk]:
        """Convert (text, heading_path) tuples to StructuredChunk list."""
        result = []
        for i, (text, heading) in enumerate(chunks):
            text = text.strip()
            if not text:
                continue
            result.append(StructuredChunk(
                id=f"chunk-{i}",
                text=text[:5000],
                heading_path=heading,
                chunk_index=i,
            ))
        return result

    def _parse_via_registry(self, file_path: str, expected_ext: str = "") -> list:
        """Parse a file via the ConverterRegistry (System A). Returns DocumentElement[] or empty list."""
        try:
            from core.harness.document.protocol import get_document_registry, StreamInfo
            import os
            ext = expected_ext or os.path.splitext(file_path)[1].lower()
            registry = get_document_registry()
            info = StreamInfo(local_path=file_path, extension=ext)
            converter = registry.find_converter(info)
            if converter:
                with open(file_path, "rb") as f:
                    return converter.convert(f, info)
            return []
        except Exception:
            return []

    def _build_chunks_from_elements(self, elements: list) -> List[StructuredChunk]:
        """Convert DocumentElement[] to StructuredChunk[] with heading path tracking.
        
        Detects markdown headings in element text (from markitdown output)
        and builds heading_path for each chunk.
        """
        full_text = "\n\n".join(el.text for el in elements if el.text.strip())
        return self._build_chunks_from_raw(full_text)

    def _build_chunks_from_raw(self, text: str) -> List[StructuredChunk]:
        """Parse raw text into heading-structured chunks. Detects markdown headings."""
        chunks = self._split_by_headings(text)
        return self._build_chunks(chunks)

    def _attach_tables_to_chunks(self, doc: ParsedDocument) -> None:
        """Attach extracted tables to their parent chunks by page number."""
        if not doc.tables or not doc.chunks:
            return
        # Build page→chunks index
        page_chunks: dict = {}
        for i, ch in enumerate(doc.chunks):
            page_chunks.setdefault(ch.page_num, []).append(i)
        for table in doc.tables:
            chunk_indices = page_chunks.get(table.page_num, [0])
            for ci in chunk_indices:
                existing = doc.chunks[ci].metadata.get("tables", [])
                existing.append(table.to_dict())
                doc.chunks[ci].metadata["tables"] = existing
                # Set chunk_id on table
                table.chunk_id = doc.chunks[ci].id

    def parse_qa_pairs(self, text: str) -> List[QAPair]:
        """Extract QA pairs from FAQ-style text (Q: ... A: ... or ### heading)."""
        import re
        pairs = []
        qa_blocks = re.split(r'\n(?=Q\s*[：:])', text)
        for i, block in enumerate(qa_blocks):
            q_match = re.match(r'Q\s*[：:]\s*(.+?)(?:\n|$)', block, re.DOTALL)
            a_match = re.search(r'A\s*[：:]\s*(.+?)(?:\n(?=Q\s*[：:])|$)', block, re.DOTALL)
            if q_match:
                question = q_match.group(1).strip()[:200]
                answer = a_match.group(1).strip()[:2000] if a_match else block[q_match.end():].strip()[:2000]
                if question and answer:
                    pairs.append(QAPair(id=f"qa-{i}", question=question, answer=answer))
        if not pairs:
            heading_pattern = re.compile(r'^#{1,3}\s+(.+)$', re.MULTILINE)
            headings = list(heading_pattern.finditer(text))
            for j, m in enumerate(headings):
                question = m.group(1).strip()[:200]
                if question.startswith(('Q:', 'A:', '第', '一、')): continue
                start = m.end()
                end = headings[j+1].start() if j+1 < len(headings) else len(text)
                answer = text[start:end].strip()[:2000]
                if question and answer:
                    pairs.append(QAPair(id=f"qa-h{j}", question=question, answer=answer))
        return pairs

    def parse_image(self, file_path: str) -> Dict[str, Any]:
        """Multi-modal: extract image metadata and OCR text via InfraOCRAdapter."""
        result = {"file": file_path, "format": "image", "text": "", "metadata": {}}
        try:
            from PIL import Image
            with Image.open(file_path) as img:
                result["metadata"] = {
                    "width": img.width, "height": img.height,
                    "mode": img.mode, "format_name": img.format,
                }
        except ImportError:
            result["text"] = f"[Image: {file_path} — PIL not available]"
            return result

        try:
            from core.harness.infrastructure.infra_ocr_adapter import create_infra_ocr_adapter
            ocr = create_infra_ocr_adapter()
            result["text"] = ocr.ocr_text(image_path=file_path)[:2000]
            result["ocr_available"] = True
        except Exception:
            result["ocr_available"] = False
            meta = result.get("metadata", {})
            result["text"] = f"[Image: {meta.get('width', '?')}x{meta.get('height', '?')} {meta.get('mode', '?')}]"
        except Exception as e:
            result["error"] = str(e)
        return result

    def _parse_video(self, file_path: str) -> ParsedDocument:
        """Parse video file via ConverterRegistry (delegates to VideoConverter).
        
        The VideoConverter handles ffmpeg audio extraction + Whisper transcription.
        System B wraps the result in ParsedDocument format.
        """
        doc = ParsedDocument(format="video", raw_text="")
        doc.title = _Path(file_path).stem

        # Delegate to registry
        elements = self._parse_via_registry(file_path, ".mp4")
        if elements:
            text = " ".join(el.text for el in elements if el.text.strip() and not el.text.startswith("["))
            if text:
                return self._parse_text(text, _Path(file_path).stem)
            else:
                doc.parse_warnings.append("No transcription produced")
                return doc

        # Fallback: direct transcriber call
        try:
            from core.harness.document.video import probe_duration_ms, extract_audio
            from core.harness.document.transcriber import transcribe_audio
            import tempfile
            duration = probe_duration_ms(file_path)
            doc.metadata["duration_ms"] = duration
            with tempfile.NamedTemporaryFile(suffix=".wav", delete=False) as tmp:
                extract_audio(file_path, tmp.name)
                segments = transcribe_audio(tmp.name)
                text = " ".join(s.get("text", "") for s in segments if s.get("text"))
                doc.metadata["transcription_segments"] = len(segments)
                if text:
                    return self._parse_text(text, _Path(file_path).stem)
        except ImportError as e:
            doc.parse_warnings.append(f"Video parsing unavailable: {e}")
        except Exception as e:
            doc.parse_warnings.append(f"Video parse error: {e}")

        doc.parse_warnings.append("No transcription produced")
        return doc

    def _parse_audio_file(self, file_path: str) -> ParsedDocument:
        """Parse audio file via ConverterRegistry (delegates to AudioConverter)."""
        doc = ParsedDocument(format="audio", raw_text="")
        doc.title = _Path(file_path).stem

        # Delegate to registry
        elements = self._parse_via_registry(file_path, ".mp3")
        if elements:
            text = " ".join(el.text for el in elements if el.text.strip() and not el.text.startswith("["))
            if text:
                return self._parse_text(text, _Path(file_path).stem)
            else:
                doc.parse_warnings.append("No transcription produced")
                return doc

        # Fallback: direct transcriber call
        try:
            from core.harness.document.transcriber import transcribe_audio
            segments = transcribe_audio(file_path)
            text = " ".join(s.get("text", "") for s in segments if s.get("text"))
            doc.metadata["transcription_segments"] = len(segments)
            if text:
                return self._parse_text(text, _Path(file_path).stem)
        except ImportError as e:
            doc.parse_warnings.append(f"Audio parsing unavailable: {e}")
        except Exception as e:
            doc.parse_warnings.append(f"Audio parse error: {e}")

        return doc

    def _parse_image_file(self, file_path: str) -> ParsedDocument:
        """Parse image file via OCR/VLM → text chunks."""
        doc = ParsedDocument(format="image", raw_text="")
        doc.title = _Path(file_path).stem
        result = self.parse_image(file_path)
        text = result.get("text", "")
        if text and not text.startswith("[Image:"):
            return self._parse_text(text, _Path(file_path).stem)
        else:
            doc.chunks = [StructuredChunk(
                id="img-0", text=text[:2000],
                heading_path=[_Path(file_path).stem],
                chunk_index=0,
            )]
            doc.metadata = result.get("metadata", {})
            return doc

    def _extract_title_md(self, text: str) -> str:
        """Extract title from first H1 in Markdown."""
        for line in text.split("\n"):
            m = _re.match(r"^#\s+(.+)", line)
            if m:
                return m.group(1).strip()[:120]
        # Fallback: first non-empty line
        for line in text.split("\n"):
            if line.strip():
                return line.strip()[:120]
        return ""
